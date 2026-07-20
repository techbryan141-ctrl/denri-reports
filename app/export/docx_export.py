"""Renders the same report_data dict served by /api/report into a .docx file.

Reuses the web dashboard's already-computed, already-formatted data verbatim -
no metric is recalculated here. One generic table renderer covers every
section's tables (they're all shaped like lists of dicts with a handful of
text fields plus an optional 'dir'/'status_tone' key for coloring), so adding
a 10th table doesn't mean writing a 10th bespoke function.
"""
import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# Palette mirrors the web dashboard's validated light-mode tokens, so the
# exported document reads as the same product, not a different one.
GOOD = RGBColor(0x0C, 0xA3, 0x0C)
WARNING = RGBColor(0xB8, 0x86, 0x00)
CRITICAL = RGBColor(0xD0, 0x3B, 0x3B)
NEUTRAL = RGBColor(0x52, 0x51, 0x4E)
MUTED = RGBColor(0x89, 0x87, 0x81)
BLUE = RGBColor(0x2A, 0x78, 0xD6)

HEADER_FILL = "1F3864"  # section headings / banner
TABLE_HEADER_FILL = "DCE3EE"  # light, never solid black
CALLOUT_FILLS = {
    "risk": "FBE4E2", "critical": "FBE4E2",
    "positive": "E4F3E4", "good": "E4F3E4", "win": "E4F3E4",
    "watch": "FDF1DA", "warning": "FDF1DA", "high": "FDF1DA", "medium": "FDF1DA",
    "key": "E3EBF8", "context": "E3EBF8", "neutral": "E3EBF8",
    "escalation": "F6D0CC",
    "meeting": "EFEFEF",
}

_DIR_COLOR = {"up": GOOD, "down": CRITICAL, "neutral": NEUTRAL}
_TONE_COLOR = {"good": GOOD, "warning": WARNING, "critical": CRITICAL, "neutral": NEUTRAL}
_PRIORITY_COLOR = {"CRITICAL": CRITICAL, "HIGH": WARNING, "MEDIUM": NEUTRAL, "WIN": GOOD}

FLAG_ICON = {"risk": "\U0001F534", "positive": "\U0001F7E2", "watch": "⚠", "key": "\U0001F511",
             "escalation": "\U0001F6A8", "context": "\U0001F5D3"}

PERIOD_TITLES = {"day": "DAILY", "week": "WEEKLY", "month": "MONTHLY"}


def _shade_paragraph(paragraph, hex_fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    pPr.append(shd)


def _shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _style_run(run, size=10, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _para(doc, text="", size=10, bold=False, italic=False, color=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        _style_run(p.add_run(text), size, bold, italic, color)
    return p


def _banner(doc, meta):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _shade_paragraph(p, HEADER_FILL)
    title = f"{PERIOD_TITLES.get(meta['period_type'], '')} SALES PERFORMANCE REPORT"
    _style_run(p.add_run(title), size=15, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    _shade_paragraph(p2, HEADER_FILL)
    _style_run(p2.add_run(f"{meta['period_label']} | {meta['date_range']} ({meta['days']} day(s))"),
               size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(2)
    _shade_paragraph(p3, HEADER_FILL)
    _style_run(p3.add_run("Denri Africa — Business Intelligence"), size=9, color=RGBColor(0xE0, 0xE0, 0xE0))

    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(10)
    _shade_paragraph(p4, HEADER_FILL)
    _style_run(p4.add_run(f"Compared to: {meta['compared_to']}"), size=9, italic=True, color=RGBColor(0xE0, 0xE0, 0xE0))


def _kpi_strip(doc, cards):
    table = doc.add_table(rows=2, cols=len(cards))
    table.autofit = True
    for i, card in enumerate(cards):
        val_cell, sub_cell = table.rows[0].cells[i], table.rows[1].cells[i]
        _shade_cell(val_cell, TABLE_HEADER_FILL)
        _shade_cell(sub_cell, TABLE_HEADER_FILL)

        vp = val_cell.paragraphs[0]
        vp.add_run(card["label"]).font.size = Pt(8)
        vp.runs[0].font.color.rgb = MUTED
        vp2 = val_cell.add_paragraph()
        _style_run(vp2.add_run(card["value"]), size=13, bold=True)

        sp = sub_cell.paragraphs[0]
        _style_run(sp.add_run(card["sub"]), size=8, bold=True, color=_TONE_COLOR.get(card["tone"], NEUTRAL))
    _para(doc, "", space_after=8)


def _callout(doc, text, kind="context", icon=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    _shade_paragraph(p, CALLOUT_FILLS.get(kind, "F2F2F2"))
    prefix = f"{icon} " if icon else ""
    _style_run(p.add_run(prefix + text), size=9.5, italic=(kind == "meeting"))


def _section_heading(doc, number, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    _style_run(p.add_run(f"{number}. {title}"), size=13, bold=True, color=RGBColor(*bytes.fromhex(HEADER_FILL)))


def _subheading(doc, number, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    _style_run(p.add_run(f"{number} {title}"), size=11, bold=True)


def _add_table(doc, headers, rows, getters, color_specs=None, widths=None):
    """getters: list of (row)->str, one per header. color_specs: optional dict
    {col_index: (row)->RGBColor|None} for tinting a column's text."""
    if not rows:
        _para(doc, "No data for this period.", size=9, italic=True, color=MUTED)
        return

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _shade_cell(hdr[i], TABLE_HEADER_FILL)
        p = hdr[i].paragraphs[0]
        _style_run(p.add_run(h), size=9, bold=True)

    for row in rows:
        cells = table.add_row().cells
        is_total = str(row.get("shop", row.get("metric", ""))).strip().upper() == "TOTAL"
        for i, getter in enumerate(getters):
            text = str(getter(row))
            p = cells[i].paragraphs[0]
            color = None
            if color_specs and i in color_specs:
                color = color_specs[i](row)
            _style_run(p.add_run(text), size=9, bold=is_total, color=color)

    if widths:
        for row_ in table.rows:
            for i, w in enumerate(widths):
                row_.cells[i].width = Cm(w)


def _dir_color(row):
    return _DIR_COLOR.get(row.get("dir"), None)


def _tone_color(row):
    return _TONE_COLOR.get(row.get("status_tone"), None)


# --- Sections ---

def _exec_summary(doc, data):
    _section_heading(doc, 1, "Executive Summary")
    _para(doc, data["narrative"], size=10)
    _add_table(
        doc, ["Metric", "Current", "Previous", "Change"], data["metrics"],
        [lambda r: r["metric"], lambda r: r["current"], lambda r: r["previous"], lambda r: r["change"]],
        color_specs={3: _dir_color},
    )
    _para(doc, "", space_after=4)
    for flag in data["flags"]:
        _callout(doc, flag["text"], kind=flag["kind"], icon=FLAG_ICON.get(flag["kind"]))
    _callout(doc, data["meeting_note"], kind="meeting", icon="\U0001F4CB")


def _customer_metrics(doc, data):
    _section_heading(doc, 2, "Customer Metrics")
    _para(doc, data["summary"], size=10)
    _add_table(
        doc, ["Metric", "Current", "Previous", "Change"], data["comparison"],
        [lambda r: r["metric"], lambda r: r["current"], lambda r: r["previous"], lambda r: r["change"]],
        color_specs={3: _dir_color},
    )
    _subheading(doc, "2.1", "Purchase Frequency Distribution")
    _add_table(
        doc, ["Purchase Frequency", "Customer Count", "Percentage", "Read"], data["frequency"],
        [lambda r: r["band"], lambda r: r["count"], lambda r: r["pct"], lambda r: r["read"]],
    )
    _subheading(doc, "2.2", "Channel Overview")
    _add_table(
        doc, ["Channel", "Orders", "Share", "Avg Spend", "Note"], data["channels"],
        [lambda r: r["channel"], lambda r: r["orders"], lambda r: r["share"], lambda r: r["avg_spend"], lambda r: r["note"]],
    )
    _callout(doc, data["meeting_note"], kind="meeting", icon="\U0001F4CB")


def _store_performance(doc, data):
    _section_heading(doc, 3, "Store Performance")
    _para(doc, data["summary"], size=10)
    _add_table(
        doc, ["Shop", "Current", "Previous", "Change", "Change %", "New", "Repeat"], data["rows"],
        [lambda r: r["shop"], lambda r: r["current"], lambda r: r["previous"], lambda r: r["change"],
         lambda r: r["change_pct"], lambda r: r["new"], lambda r: r["repeat"]],
        color_specs={3: _dir_color, 4: _dir_color},
    )
    _subheading(doc, "3.1", "Channel Mix by Location")
    _add_table(
        doc, ["Shop", "Walk-in", "Online", "Activation", "Total", "Online %"], data["channel_mix"],
        [lambda r: r["shop"], lambda r: r["walkin"], lambda r: r["online"], lambda r: r["activation"],
         lambda r: r["total"], lambda r: r["online_pct"]],
    )
    _callout(doc, data["meeting_note"], kind="meeting", icon="\U0001F4CB")


def _gender_performance(doc, data):
    _section_heading(doc, 4, "Gender Performance Analysis")
    _para(doc, data["summary"], size=10)
    if data.get("ratio"):
        _para(doc, f"Ratio: {data['ratio']['ratio_text']}", size=9, bold=True, color=BLUE)

    _subheading(doc, "4.1", "Overall Gender Split")
    _add_table(
        doc, ["Gender", "Count", "%", "Change"], data["overall"],
        [lambda r: r["gender"], lambda r: r["count"], lambda r: r["pct"], lambda r: r["change"]],
        color_specs={3: _dir_color},
    )

    _subheading(doc, "4.2", "Gender by Location")
    cols = data["by_location"]["columns"]
    headers = ["Shop"] + cols + ["Total"]
    getters = [lambda r: r["shop"]] + [(lambda r, c=c: r[c.lower()]) for c in cols] + [lambda r: r["total"]]
    _add_table(doc, headers, data["by_location"]["rows"], getters)

    _callout(doc, data["meeting_note"], kind="meeting", icon="\U0001F4CB")


def _traffic(doc, data):
    _section_heading(doc, 5, "Foot & Online Traffic Analysis")
    _para(doc, data["summary"], size=10)
    if data.get("data_gap_note"):
        _callout(doc, data["data_gap_note"], kind="watch", icon="⚠")
    _add_table(
        doc, ["Shop", "Walk-in Purchased", "Walk-in Total", "Conv. Rate", "Online", "Activation", "Total Customers"],
        data["rows"],
        [lambda r: r["shop"], lambda r: r["walkin_purchased"], lambda r: r["walkin_total"], lambda r: r["conv_rate"],
         lambda r: r["online"], lambda r: r["activation"], lambda r: r["total_customers"]],
    )
    _callout(doc, data["meeting_note"], kind="meeting", icon="\U0001F4CB")


def _revenue_analysis(doc, data):
    _section_heading(doc, 6, "Revenue Analysis")
    _para(doc, data["summary"], size=10)
    _kpi_strip(doc, data["kpi_cards"])
    _add_table(
        doc, ["Metric", "Current", "Previous", "Change"], data["comparison"],
        [lambda r: r["metric"], lambda r: r["current"], lambda r: r["previous"], lambda r: r["change"]],
        color_specs={3: _dir_color},
    )
    _subheading(doc, "6.1", "Revenue Bridge")
    _add_table(
        doc, ["Period", "Revenue", "Customers", "Avg Spend", "Trend"], data["bridge"],
        [lambda r: r["period"], lambda r: r["revenue"], lambda r: r["customers"], lambda r: r["avg_spend"], lambda r: r["trend"]],
        color_specs={4: _dir_color},
    )
    _callout(doc, data["meeting_note"], kind="meeting", icon="\U0001F4CB")


def _data_quality(doc, data):
    _section_heading(doc, 7, "Data Quality (KYC)")
    _para(doc, data["summary"], size=10)
    _add_table(
        doc, ["Metric", "Current", "Current %", "Previous", "Previous %", "Change"], data["comparison"],
        [lambda r: r["metric"], lambda r: r["current"], lambda r: r["current_pct"], lambda r: r["previous"],
         lambda r: r["previous_pct"], lambda r: r["change"]],
        color_specs={5: _dir_color},
    )
    if data.get("store_number_by_shop"):
        _subheading(doc, "7.1", "Store Numbers Used by Location")
        _add_table(
            doc, ["Shop", "Unique Numbers", "Occurrences"], data["store_number_by_shop"],
            [lambda r: r["shop"], lambda r: r["unique_numbers"], lambda r: r["occurrences"]],
        )
    _callout(doc, data["score_note"], kind="key", icon="\U0001F511")
    _callout(doc, data["meeting_note"], kind="meeting", icon="\U0001F4CB")


def _feedback(doc, data):
    _section_heading(doc, 8, "Customer Feedback")
    _para(doc, data["summary"], size=10)
    _add_table(
        doc, ["Metric", "Current", "Previous", "Change %", "Status"], data["comparison"],
        [lambda r: r["metric"], lambda r: r["current"], lambda r: r["previous"], lambda r: r["change"], lambda r: r["status"]],
        color_specs={3: _dir_color, 4: _tone_color},
    )
    if data.get("target"):
        _callout(doc, data["target"]["note"], kind=data["target"].get("status_tone", "context"), icon="\U0001F3AF")

    _subheading(doc, "8.1", "Store Breakdown")
    _add_table(
        doc, ["Shop", "Current", "Previous", "Change"], data["store_breakdown"],
        [lambda r: r["shop"], lambda r: r["current"], lambda r: r["previous"], lambda r: r["change"]],
        color_specs={3: _dir_color},
    )

    if data.get("feedback_links"):
        _subheading(doc, "8.2", "Feedback Link Targets by Location")
        _add_table(
            doc, ["Shop", "Links Sent", "Online", "Achv %", "Gap", "Status", "Walk-Ins"], data["feedback_links"],
            [lambda r: r["shop"], lambda r: r["links_sent"], lambda r: r["online"], lambda r: r["achv_pct"],
             lambda r: r["gap"], lambda r: r["status"], lambda r: r["walkins"]],
            color_specs={5: _tone_color},
        )

    _callout(doc, data["meeting_note"], kind="meeting", icon="\U0001F4CB")


def _priorities(doc, data):
    _section_heading(doc, 9, "Summary & Key Notes")
    dept_labels = {"marketing": "Marketing", "sales": "Sales", "production": "Production", "data": "Data"}
    for key, label in dept_labels.items():
        dept = data[key]
        _subheading(doc, "•", label)
        _para(doc, dept["summary"], size=9.5, italic=True, color=MUTED)
        if dept.get("unavailable"):
            _para(doc, dept["note"], size=9, italic=True, color=MUTED)
            continue
        for item in dept["items"]:
            p = doc.add_paragraph(style=None)
            p.paragraph_format.space_after = Pt(4)
            _style_run(p.add_run(f"[{item['priority']}] "), size=9, bold=True, color=_PRIORITY_COLOR.get(item["priority"]))
            _style_run(p.add_run(item["text"]), size=9)


def _footer(doc, meta):
    _para(doc, "", space_after=10)
    p = doc.add_paragraph()
    _style_run(
        p.add_run(f"Generated by: Denri Africa Reporting Dashboard | {meta['generated_on']} | Denri Africa — Business Intelligence"),
        size=8, italic=True, color=MUTED,
    )


def build_docx(report: dict) -> io.BytesIO:
    doc = Document()
    section = doc.sections[0]
    section.page_height, section.page_width = Cm(29.7), Cm(21.0)  # A4 portrait
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Cm(1.9))

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    _banner(doc, report["meta"])
    _kpi_strip(doc, report["kpi_strip"])
    _callout(doc, report["context_note"], kind="context", icon="\U0001F5D3")

    _exec_summary(doc, report["exec_summary"])
    _customer_metrics(doc, report["customer_metrics"])
    _store_performance(doc, report["store_performance"])
    _gender_performance(doc, report["gender_performance"])
    _traffic(doc, report["traffic"])
    _revenue_analysis(doc, report["revenue_analysis"])
    _data_quality(doc, report["data_quality"])
    _feedback(doc, report["feedback"])
    _priorities(doc, report["priorities"])
    _footer(doc, report["meta"])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
